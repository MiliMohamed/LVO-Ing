# -*- coding: utf-8 -*-
"""
Génération du compte-rendu trimestriel Word (docxtpl).
"""

from __future__ import annotations

import io
from datetime import date
from pathlib import Path
from typing import Any, Optional

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches
from docxtpl import DocxTemplate

from config import parametres as cfg

TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"
ASSETS_DIR = Path(__file__).resolve().parent.parent / "assets"
MODELE_CR = TEMPLATES_DIR / "CR_modele.docx"
LOGO_SVG = ASSETS_DIR / "logo-lvo.svg"
LOGO_COLOR_PNG = ASSETS_DIR / "logo_lvo_color.png"


def creer_modele_si_absent() -> Path:
    """
    Crée le modèle Word avec balises Jinja2 si le fichier n'existe pas encore.
    """
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    if MODELE_CR.exists():
        return MODELE_CR

    doc = Document()
    doc.add_heading("Compte-rendu trimestriel — Maintenance ascenseurs", level=0)
    doc.add_paragraph("Prestataire : {{ prestataire }}")
    doc.add_paragraph("Période : {{ trimestre }} {{ annee }}")
    doc.add_paragraph("Date du compte-rendu : {{ date_cr }}")

    doc.add_heading("1. Rappel du parc", level=1)
    doc.add_paragraph("Nombre d'appareils suivis : {{ nb_appareils }}")
    doc.add_paragraph("{{ liste_parc }}")

    doc.add_heading("2. Visites d'entretien", level=1)
    doc.add_paragraph("Nombre de visites : {{ nb_visites_maintenance }}")
    doc.add_paragraph("Écarts supérieurs à {{ seuil_maintenance }} jours : {{ nb_ecarts_maintenance }}")
    doc.add_paragraph("Jours de retard cumulés : {{ jours_retard_maintenance }}")
    doc.add_paragraph("Pénalité maintenance : {{ penalite_maintenance }} €")

    doc.add_heading("3. Tests parachute et contrôles câbles", level=1)
    doc.add_paragraph("Tests parachute réalisés : {{ nb_parachute }}")
    doc.add_paragraph("Contrôles câbles réalisés : {{ nb_cables }}")
    doc.add_paragraph("Appareils sans test parachute : {{ appareils_sans_parachute }}")
    doc.add_paragraph("Appareils sans contrôle câbles : {{ appareils_sans_cables }}")

    doc.add_heading("4. Fonctionnement / taux de pannes", level=1)
    doc.add_paragraph("Nombre total d'interventions : {{ nb_interventions }}")
    doc.add_paragraph("Pannes retenues (imputables) : {{ nb_pannes_retenues }}")
    doc.add_paragraph("Appareils en excès de pannes : {{ appareils_exces_pannes }}")
    doc.add_paragraph("Pénalité pannes : {{ penalite_pannes }} €")

    doc.add_heading("5. Délais d'intervention", level=1)
    doc.add_paragraph("Répartition :")
    doc.add_paragraph("  • Moins de 1 h : {{ delai_moins_1h }}")
    doc.add_paragraph("  • 1 à 2 h : {{ delai_1_2h }}")
    doc.add_paragraph("  • 2 à 4 h : {{ delai_2_4h }}")
    doc.add_paragraph("  • Plus de 4 h : {{ delai_plus_4h }}")
    doc.add_paragraph("Pénalité délais : {{ penalite_delai }} €")

    doc.add_heading("6. Désincarcération", level=1)
    doc.add_paragraph("Nombre d'interventions avec personne bloquée : {{ nb_desincarceration }}")
    doc.add_paragraph("Interventions hors délai contractuel ({{ seuil_desinc }} min) : {{ nb_hors_delai_desinc }}")

    doc.add_heading("7. Synthèse des pénalités", level=1)
    doc.add_paragraph("Pénalité maintenance : {{ penalite_maintenance }} €")
    doc.add_paragraph("Pénalité pannes : {{ penalite_pannes }} €")
    doc.add_paragraph("Pénalité immobilisation : {{ penalite_immobilisation }} €")
    doc.add_paragraph("Pénalité délais d'intervention : {{ penalite_delai }} €")
    doc.add_paragraph("Montant total des pénalités : {{ penalite_totale }} €")
    doc.add_paragraph("")
    doc.add_paragraph("Commentaire : {{ commentaire }}")

    doc.save(str(MODELE_CR))
    return MODELE_CR


def _contexte_word(
    resultats: dict[str, Any],
    params: dict[str, Any],
) -> dict[str, Any]:
    """Construit le dictionnaire de variables pour docxtpl."""
    interventions = resultats.get("interventions", pd.DataFrame())
    maintenance = resultats.get("maintenance", pd.DataFrame())
    suivi = resultats.get("suivi_tests", pd.DataFrame())
    synthese = resultats.get("synthese_appareils", pd.DataFrame())
    pen = resultats.get("penalites", {})
    detail = pen.get("detail", {})
    desinc = resultats.get("desincarceration", pd.DataFrame())
    parc_orphelins = resultats.get("parc_orphelins", pd.DataFrame())

    nb_appareils = len(synthese) if not synthese.empty else 0
    liste_parc = ""
    if not synthese.empty:
        lignes = [
            f"- {r.get('client', '')} / {r.get('ascenseur', '')} — {r.get('adresse', '')}"
            for _, r in synthese.iterrows()
        ]
        liste_parc = "\n".join(lignes[:50])
        if len(lignes) > 50:
            liste_parc += f"\n… et {len(lignes) - 50} autres appareils."

    tranches = interventions["tranche_delai"].value_counts().to_dict() if not interventions.empty and "tranche_delai" in interventions.columns else {}

    app_exces = 0
    if not synthese.empty:
        from core.calculs import seuil_prorata_trimestre

        seuil = seuil_prorata_trimestre(
            float(params.get("seuil_pannes_an", cfg.SEUIL_PANNES_PAR_AN_APPAREIL)),
            str(params.get("trimestre", "T1")),
        )
        app_exces = int((synthese["nb_pannes_retenues"] > seuil).sum())

    sans_para = int(suivi["parachute_manquant"].sum()) if not suivi.empty and "parachute_manquant" in suivi.columns else 0
    sans_cables = int(suivi["cables_manquant"].sum()) if not suivi.empty and "cables_manquant" in suivi.columns else 0

    nb_ecarts = int(maintenance["ecart_signale"].sum()) if not maintenance.empty and "ecart_signale" in maintenance.columns else 0

    return {
        "prestataire": params.get("prestataire", "—"),
        "trimestre": params.get("trimestre", "T1"),
        "annee": params.get("annee", date.today().year),
        "date_cr": params.get("date_cr", date.today().strftime("%d/%m/%Y")),
        "nb_appareils": nb_appareils,
        "liste_parc": liste_parc or "Non renseigné",
        "nb_visites_maintenance": len(maintenance),
        "seuil_maintenance": params.get("seuil_maintenance_jours", cfg.SEUIL_ECART_MAINTENANCE_JOURS),
        "nb_ecarts_maintenance": nb_ecarts,
        "jours_retard_maintenance": detail.get("jours_retard_maintenance", 0),
        "penalite_maintenance": pen.get("penalite_maintenance", 0),
        "nb_parachute": len(resultats.get("parachute", pd.DataFrame())),
        "nb_cables": len(resultats.get("cables", pd.DataFrame())),
        "appareils_sans_parachute": sans_para,
        "appareils_sans_cables": sans_cables,
        "nb_interventions": len(interventions),
        "nb_pannes_retenues": int(interventions["panne_retenue"].sum()) if not interventions.empty and "panne_retenue" in interventions.columns else 0,
        "appareils_exces_pannes": app_exces,
        "penalite_pannes": pen.get("penalite_pannes", 0),
        "delai_moins_1h": tranches.get("< 1h", 0),
        "delai_1_2h": tranches.get("1 à 2h", 0),
        "delai_2_4h": tranches.get("2 à 4h", 0),
        "delai_plus_4h": tranches.get("> 4h", 0),
        "penalite_delai": pen.get("penalite_delai_intervention", 0),
        "nb_desincarceration": len(desinc),
        "nb_hors_delai_desinc": int(desinc["hors_delai_45min"].sum()) if not desinc.empty and "hors_delai_45min" in desinc.columns else 0,
        "seuil_desinc": params.get("seuil_desincarc_min", cfg.SEUIL_DESINCARCERATION_MINUTES),
        "penalite_immobilisation": pen.get("penalite_immobilisation", 0),
        "penalite_totale": pen.get("penalite_totale", 0),
        "commentaire": params.get("commentaire", ""),
        "libelle_mode_pannes": detail.get("libelle_mode_pannes", ""),
        "libelle_mode_immobilisation": detail.get("libelle_mode_immobilisation", ""),
        "mention_pannes": detail.get("mention_pannes", ""),
        "client": params.get("filtre_client") or params.get("hypotheses_client") or params.get("client_site", ""),
    }


def _assurer_logo_couleur() -> Optional[Path]:
    """Génère un PNG charte LVO si absent (logo couleur pour les rapports)."""
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    if LOGO_COLOR_PNG.is_file():
        return LOGO_COLOR_PNG
    if LOGO_SVG.is_file():
        try:
            import cairosvg

            cairosvg.svg2png(url=str(LOGO_SVG), write_to=str(LOGO_COLOR_PNG), output_width=800)
            return LOGO_COLOR_PNG
        except Exception:
            pass
    try:
        import matplotlib.pyplot as plt

        fig, ax = plt.subplots(figsize=(4, 1.2))
        ax.set_axis_off()
        ax.text(
            0.5,
            0.55,
            "LVO",
            fontsize=42,
            fontweight="bold",
            color="#1F3A5F",
            ha="center",
            va="center",
            transform=ax.transAxes,
        )
        ax.text(
            0.5,
            0.15,
            "Ingénierie",
            fontsize=14,
            color="#E67E22",
            ha="center",
            va="center",
            transform=ax.transAxes,
        )
        fig.savefig(LOGO_COLOR_PNG, dpi=200, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        return LOGO_COLOR_PNG
    except Exception:
        return None


def _enrichir_document_logo_et_graphiques(
    docx_bytes: bytes,
    figure_paths: Optional[dict[str, str]] = None,
) -> bytes:
    """Ajoute le logo LVO en en-tête et les figures en fin de document."""
    buffer_in = io.BytesIO(docx_bytes)
    doc = Document(buffer_in)

    logo = _assurer_logo_couleur()
    if logo and logo.is_file():
        if doc.sections:
            header = doc.sections[0].header
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run()
            run.add_picture(str(logo), width=Cm(4.5))

    if figure_paths:
        doc.add_page_break()
        doc.add_heading("Graphiques", level=1)
        for cle, chemin in figure_paths.items():
            p = Path(chemin)
            if not p.is_file():
                continue
            doc.add_paragraph(cle.replace("_", " ").title())
            try:
                doc.add_picture(str(p), width=Inches(6.0))
            except Exception:
                continue

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


def generer_compte_rendu(
    resultats: dict[str, Any],
    params: dict[str, Any],
    figure_paths: Optional[dict[str, str]] = None,
) -> bytes:
    """Génère le .docx en mémoire (logo LVO + graphiques intégrés)."""
    creer_modele_si_absent()
    tpl = DocxTemplate(str(MODELE_CR))
    contexte = _contexte_word(resultats, params)
    tpl.render(contexte)
    buffer = io.BytesIO()
    tpl.save(buffer)
    buffer.seek(0)
    return _enrichir_document_logo_et_graphiques(buffer.getvalue(), figure_paths)


def nom_fichier_cr(params: dict[str, Any], version: int = 1) -> str:
    """Convention LVO_<CLIENT>_<PRESTATAIRE>_CR_T<n>_<année>_v<version>.docx"""
    client = str(
        params.get("filtre_client")
        or params.get("hypotheses_client")
        or params.get("client_site")
        or "CLIENT"
    ).replace(" ", "_")
    if client.upper() == "TOUS":
        client = str(params.get("client_site", "CLIENT")).replace(" ", "_")
    prest = str(params.get("prestataire", "PRESTATAIRE")).replace(" ", "_")
    trimestre = str(params.get("trimestre", "T1"))
    annee = int(params.get("annee", date.today().year))
    return f"LVO_{client}_{prest}_CR_{trimestre}_{annee}_v{version}.docx"
