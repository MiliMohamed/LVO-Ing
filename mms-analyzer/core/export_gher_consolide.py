# -*- coding: utf-8 -*-
"""Rapport Word consolidé GHER — synthèse multi-prestataires."""

from __future__ import annotations

import io
from datetime import date
from pathlib import Path
from typing import Any, Optional

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor

from core.export_word import _assurer_logo_couleur, _enrichir_document_logo_et_graphiques
from core.figures import CHARTE


def _fig_barres_prestataires(consolide: list[dict[str, Any]]) -> Path | None:
    if not consolide:
        return None
    labels = [f"{r.get('prestataire', '?')}" for r in consolide]
    vals = [float(r.get("penalite_totale", 0)) for r in consolide]
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.bar(labels, vals, color=CHARTE["navy"])
    ax.set_ylabel("Pénalités (€)")
    ax.set_title("Pénalités par prestataire — GHER")
    plt.xticks(rotation=25, ha="right")
    import os
    import tempfile

    fd, path = tempfile.mkstemp(suffix=".png")
    os.close(fd)
    p = Path(path)
    fig.savefig(p, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return p


def _fig_camembert_part(consolide: list[dict[str, Any]]) -> Path | None:
    if not consolide:
        return None
    vals = [float(r.get("penalite_totale", 0)) for r in consolide]
    if sum(vals) <= 0:
        return None
    labels = [str(r.get("prestataire", "?")) for r in consolide]
    fig, ax = plt.subplots(figsize=(6, 5))
    ax.pie(vals, labels=labels, autopct="%1.0f%%", colors=CHARTE["palette"][: len(vals)], startangle=90)
    ax.set_title("Part des pénalités par prestataire")
    import tempfile
    import os

    fd, path = tempfile.mkstemp(suffix=".png")
    os.close(fd)
    p = Path(path)
    fig.savefig(p, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return p


def generer_cr_gher_consolide(
    consolide: list[dict[str, Any]],
    params: dict[str, Any],
    couples_parc: Optional[list[dict[str, Any]]] = None,
) -> bytes:
    """Produit un DOCX de synthèse comparative pour le groupement GHER."""
    trimestre = str(params.get("trimestre", "T1"))
    annee = int(params.get("annee", date.today().year))
    date_cr = str(params.get("date_cr", date.today().strftime("%d/%m/%Y")))

    doc = Document()
    logo = _assurer_logo_couleur()
    if logo and logo.is_file():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(str(logo), width=Cm(5))

    titre = doc.add_heading("Compte-rendu consolidé — Groupement GHER", level=0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Période : {trimestre} {annee}")
    doc.add_paragraph(f"Date du rapport : {date_cr}")
    doc.add_paragraph(f"Nombre de prestataires analysés : {len(consolide)}")

    total = sum(float(r.get("penalite_totale", 0)) for r in consolide)
    p_total = doc.add_paragraph()
    run = p_total.add_run(f"Total des pénalités proposées : {total:,.2f} €")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    detail = params.get("libelle_mode_pannes") or params.get("mention_pannes")
    if detail:
        doc.add_paragraph(f"Mode pannes : {detail}")

    doc.add_heading("Tableau comparatif par prestataire", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, lbl in enumerate(
        ["Prestataire", "Maintenance €", "Pannes €", "Immo €", "Délai €", "Total €"]
    ):
        hdr[i].text = lbl

    for row in consolide:
        cells = table.add_row().cells
        cells[0].text = str(row.get("prestataire", "—"))
        cells[1].text = f"{float(row.get('penalite_maintenance', 0)):,.2f}"
        cells[2].text = f"{float(row.get('penalite_pannes', 0)):,.2f}"
        cells[3].text = f"{float(row.get('penalite_immobilisation', 0)):,.2f}"
        cells[4].text = f"{float(row.get('penalite_delai', 0)):,.2f}"
        cells[5].text = f"{float(row.get('penalite_totale', 0)):,.2f}"

    doc.add_heading("Graphiques comparatifs", level=1)
    fig_paths: list[Path] = []
    bar = _fig_barres_prestataires(consolide)
    if bar:
        fig_paths.append(bar)
        doc.add_picture(str(bar), width=Inches(6))
    pie = _fig_camembert_part(consolide)
    if pie:
        fig_paths.append(pie)
        doc.add_picture(str(pie), width=Inches(5))

    if couples_parc:
        doc.add_heading("Répartition du parc par prestataire", level=1)
        par_prest: dict[str, int] = {}
        for c in couples_parc:
            prest = str(c.get("prestataire", "INCONNU"))
            par_prest[prest] = par_prest.get(prest, 0) + int(c.get("nb_appareils", 0))
        for prest, nb in sorted(par_prest.items()):
            doc.add_paragraph(f"• {prest} : {nb} appareil(s)", style="List Bullet")

    doc.add_heading("Détail par prestataire", level=1)
    doc.add_paragraph(
        "Les rapports individuels (Excel + Word + PDF) sont disponibles pour chaque couple "
        "client × prestataire traité dans ce lot."
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    raw = buffer.getvalue()

    for p in fig_paths:
        try:
            p.unlink(missing_ok=True)
        except OSError:
            pass

    return _enrichir_document_logo_et_graphiques(raw, None)


def nom_fichier_gher_consolide(params: dict[str, Any], version: int = 1) -> str:
    trimestre = str(params.get("trimestre", "T1"))
    annee = int(params.get("annee", date.today().year))
    return f"LVO_GHER_CONSOLIDE_CR_{trimestre}_{annee}_v{version}.docx"
